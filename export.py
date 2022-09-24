import os
import io

import numpy as np

from docx import Document
from docx.shared import Cm

from matplotlib import pyplot as plt
from matplotlib.ticker import StrMethodFormatter

from dic_data_crack_extraction import load_data_from_file


def draw_strain_field(data, x, y, time, indicies):
    '''
    Draw normal strain field
    '''
    # plt.figure()
    fig, axes = plt.subplots(nrows=2, ncols=2, figsize=(10,7))

    for ax, index in zip(axes.flat, indicies):
        im = ax.contourf(x, y, data[index], levels=255, cmap='rainbow')
        ax.set_xlabel('$\it{x}$, мм')
        ax.set_ylabel('$\it{y}$, мм')
        ax.title.set_text(f'Время = {time[index]:.2f} с')
        ax.grid()

        cbar = plt.colorbar(im, ax=ax, format=StrMethodFormatter('{x:,.4f}'))
        cbar.ax.set_ylabel('Максимальная\n относительная деформация $\it{\epsilon}$')

    plt.subplots_adjust(left=0.07, right=0.93, bottom=0.07, top=0.93, wspace=0.4, hspace=0.3)
    # plt.show()


if __name__ == '__main__':
    
    DATA_PATH = './data'
    FILE_NAME_TO_SAVE = 'appendix.docx'

    DATA_SETS = ('R_6_28_12_2021', 'O16_12_08_2021', 'H_9_2_17_12_2021', 'R_Nemo_10_03_2022', 'H_6_3_17_12_2021', 'H_3_3_26_11_2021')
    REMOTE_DIC_DATA_LINKS = ('https://yadi.sk/d/H1vUyhigblohFA', 'https://yadi.sk/d/rn3MWI0-RfzIcg', 'https://yadi.sk/d/9jzHGOotHuD9xQ', 'https://yadi.sk/d/rMMyPa9YhAxVtw', '', '')

    LAST_RECSS = (7923, 1199, 5758, 2666, 207, 11259)
    START_TIME_INDICIES = (0, 600, 5000) 

    # Index of data set to calculate
    DATA_SET_INDEX = 1
    
    DATA_SET_NAME = DATA_SETS[DATA_SET_INDEX]
    REMOTE_DIC_DATA_LINK = REMOTE_DIC_DATA_LINKS[DATA_SET_INDEX]

    # Last record to process (before destruction)    
    LAST_REC = LAST_RECSS[DATA_SET_INDEX]

    IMAGES_NUMBER = 20
    GRAPHS_PER_IMAGE = 4
    START_TIME_INDEX = START_TIME_INDICIES[DATA_SET_INDEX]

    print(f'Dataset "{DATA_SET_NAME}" choosed to process..')

    # Load data from file
    x_coords, y_coords, time_counts, data = load_data_from_file(DATA_SET_NAME, DATA_PATH)
      
    # Cut time for last record
    time_counts = time_counts[:LAST_REC]        
    
    bufs = []
    time_indecies = np.linspace(START_TIME_INDEX, LAST_REC - 1, IMAGES_NUMBER * GRAPHS_PER_IMAGE).astype(int)

    # Plotting graphs
    print('Generating images...')

    for i in range(IMAGES_NUMBER):
        draw_strain_field(data, x_coords, y_coords, time_counts, time_indecies[i*GRAPHS_PER_IMAGE:i*GRAPHS_PER_IMAGE+4])
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        bufs.append(buf)
        print(f'Image #{i} generated')

    # Plotting graphs
    print('Insert images to Word document...')

    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    p = document.add_paragraph('Приложение')
    p.bold = True

    for buf in (bufs):
        document.add_picture(buf, width=Cm(16.5), height=Cm(11.0))

    # document.add_page_break()

    document.save(FILE_NAME_TO_SAVE)
    os.system(f'start {FILE_NAME_TO_SAVE}')

    print('Script finished')
